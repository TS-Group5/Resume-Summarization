"""ATS parser for resume files."""
import logging
import re
from datetime import datetime
from typing import Dict, Any, List, Optional
from docx import Document

from .base_parser import BaseParser

logger = logging.getLogger(__name__)


class ATSParser(BaseParser):
    """Parser for ATS resume files."""
    
    # Technical and Engineering Skills
    TECHNICAL_SKILLS = {
        # Software Development
        'python', 'java', 'javascript', 'c++', 'ruby', 'php', 'swift',
        'react', 'angular', 'vue.js', 'node.js', 'django', 'flask',
        'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform',
        'git', 'ci/cd', 'jenkins', 'agile', 'scrum', 'devops',
        
        # Data Science & Analytics
        'machine learning', 'deep learning', 'artificial intelligence',
        'data analytics', 'statistical analysis', 'python', 'r',
        'sql', 'tableau', 'power bi', 'data visualization',
        'natural language processing', 'computer vision',
        
        # Project & Product
        'project management', 'agile methodologies', 'scrum',
        'product development', 'requirements gathering',
        'stakeholder management', 'risk management',
        'quality assurance', 'software development lifecycle',
        
        # Infrastructure & Security
        'network security', 'cloud computing', 'system administration',
        'infrastructure management', 'cybersecurity', 'penetration testing',
        'vulnerability assessment', 'firewall configuration'
    }
    
    # Business and Management Skills
    BUSINESS_SKILLS = {
        # Strategy & Leadership
        'strategic planning', 'business strategy', 'team leadership',
        'change management', 'organizational development',
        'business development', 'executive leadership',
        'performance management', 'operational excellence',
        
        # Finance & Operations
        'financial analysis', 'budgeting', 'forecasting',
        'risk assessment', 'process improvement', 'cost reduction',
        'supply chain management', 'inventory management',
        'vendor management', 'contract negotiation',
        
        # Marketing & Sales
        'digital marketing', 'content strategy', 'social media marketing',
        'brand management', 'market research', 'sales strategy',
        'customer relationship management', 'lead generation',
        'email marketing', 'seo optimization',
        
        # Project Management
        'project coordination', 'resource allocation',
        'timeline management', 'budget management',
        'scope management', 'risk mitigation',
        'stakeholder communication', 'vendor coordination'
    }
    
    # Creative and Design Skills
    CREATIVE_SKILLS = {
        # Design
        'ui design', 'ux design', 'graphic design', 'web design',
        'adobe creative suite', 'photoshop', 'illustrator',
        'indesign', 'sketch', 'figma', 'typography',
        
        # Content Creation
        'content writing', 'copywriting', 'technical writing',
        'blog writing', 'content strategy', 'storytelling',
        'video production', 'animation', 'photography',
        
        # Marketing Design
        'brand design', 'marketing collateral', 'visual identity',
        'campaign design', 'social media design',
        'presentation design', 'email design',
        
        # UX Research
        'user research', 'usability testing', 'wireframing',
        'prototyping', 'information architecture',
        'user journey mapping', 'a/b testing'
    }
    
    # Healthcare and Medical Skills
    HEALTHCARE_SKILLS = {
        # Clinical
        'patient care', 'medical diagnosis', 'treatment planning',
        'clinical research', 'medical documentation', 'patient assessment',
        'healthcare compliance', 'medical terminology',
        
        # Administrative
        'healthcare management', 'medical billing', 'medical coding',
        'electronic health records', 'hipaa compliance',
        'healthcare operations', 'quality improvement',
        
        # Specialized Areas
        'emergency medicine', 'pediatrics', 'geriatrics',
        'mental health', 'rehabilitation', 'pharmacy',
        'laboratory management', 'radiology',
        
        # Research & Development
        'clinical trials', 'medical research', 'drug development',
        'patient safety', 'quality control', 'regulatory compliance'
    }
    
    # HR and People Management Skills
    HR_SKILLS = {
        # Recruitment
        'talent acquisition', 'recruitment', 'interviewing',
        'candidate sourcing', 'onboarding', 'employer branding',
        'job description writing', 'applicant tracking systems',
        
        # Employee Relations
        'employee relations', 'conflict resolution', 'performance management',
        'employee engagement', 'workplace culture', 'diversity & inclusion',
        'employee development', 'succession planning',
        
        # HR Operations
        'hr policies', 'benefits administration', 'compensation',
        'payroll management', 'hr compliance', 'labor relations',
        'hr analytics', 'hris management',
        
        # Training & Development
        'training program development', 'learning management systems',
        'leadership development', 'skill assessment',
        'career development', 'mentoring programs'
    }
    
    # Soft Skills
    SOFT_SKILLS = {
        # Communication
        'verbal communication', 'written communication',
        'presentation skills', 'public speaking', 'active listening',
        'interpersonal skills', 'negotiation', 'conflict resolution',
        
        # Leadership
        'team leadership', 'mentoring', 'coaching', 'delegation',
        'decision making', 'strategic thinking', 'problem solving',
        'change management', 'innovation',
        
        # Personal Qualities
        'adaptability', 'creativity', 'critical thinking',
        'time management', 'organization', 'attention to detail',
        'initiative', 'reliability', 'professionalism',
        
        # Collaboration
        'team collaboration', 'cross-functional coordination',
        'relationship building', 'cultural awareness',
        'remote collaboration', 'partnership management'
    }
    
    # Education and Research Skills
    EDUCATION_SKILLS = {
        # Teaching
        'curriculum development', 'lesson planning', 'student assessment',
        'classroom management', 'educational technology',
        'distance learning', 'special education',
        
        # Research
        'research methodology', 'data collection', 'data analysis',
        'academic writing', 'grant writing', 'literature review',
        'experimental design', 'statistical analysis',
        
        # Administration
        'educational leadership', 'program development',
        'student counseling', 'academic advising',
        'educational policy', 'accreditation',
        
        # Technology
        'learning management systems', 'educational software',
        'online teaching tools', 'digital assessment',
        'virtual classroom management'
    }
    
    IRRELEVANT_TERMS = {
        # Personal Information
        'gpa', 'university', 'college', 'school', 'degree',
        'graduate', 'undergraduate', 'diploma', 'certificate',
        
        # Activities
        'activities', 'hobbies', 'interests', 'volunteer',
        'club', 'society', 'association', 'member',
        
        # General Terms
        'languages', 'travel', 'sports', 'awards',
        'references', 'portfolio', 'availability',
        
        # Contact Info
        'email', 'phone', 'address', 'linkedin',
        'website', 'social media', 'github'
    }
    
    def __init__(self, file_path: str):
        """Initialize the parser with a file path."""
        super().__init__(file_path)
        self.resume_data = {
            'name': '',
            'current_role': '',
            'companies': [],
            'years_experience': 0,
            'skills': [],
            'achievements': [],
            'education': [],
            'contact_info': {}
        }
    
    def parse(self) -> Dict[str, Any]:
        """Parse the resume file and extract structured information.
        
        Returns:
            Dictionary containing parsed resume data
        """
        try:
            # Read the Word document
            doc = Document(self.file_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Process the content
            self._process_content(text_content)
            
            # Log the extracted data
            for key, value in self.resume_data.items():
                logger.info(f"{key.title()}: {value}")
            
            return self.resume_data
            
        except Exception as e:
            logger.error(f"Error parsing file {self.file_path}: {str(e)}")
            raise

    def clean_text(self, text):
        """Clean and normalize text."""
        cleaned = ' '.join(text.strip().split())
        logger.debug(f"Cleaned text: '{text}' -> '{cleaned}'")
        return cleaned

    def parse_date(self, text):
        """Parse date from text in various formats."""
        logger.debug(f"Parsing date from: '{text}'")
        
        # Remove any non-alphanumeric characters from the end
        text = text.strip().rstrip('.')
        
        # Handle 'Present' or 'Current'
        if text.lower() in ['present', 'current']:
            logger.debug("Found 'present/current', using current date")
            return datetime.now()
        
        # Handle 20XX format
        if re.match(r'20XX', text, re.IGNORECASE):
            logger.debug("Found '20XX' format, using 2021")
            return datetime(2021, 1, 1)  # Assume recent
        
        # Common date formats
        date_patterns = [
            (r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[\s,]+20\d{2}', '%B %Y'),
            (r'20\d{2}', '%Y')
        ]
        
        for pattern, date_format in date_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                try:
                    date_str = match.group(0)
                    # Standardize month abbreviations
                    date_str = re.sub(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*',
                                    lambda m: m.group(0)[:3], date_str, flags=re.IGNORECASE)
                    date_str = re.sub(r'[,\s]+', ' ', date_str).strip()
                    parsed_date = datetime.strptime(date_str, date_format)
                    logger.debug(f"Successfully parsed date: {date_str} -> {parsed_date}")
                    return parsed_date
                except ValueError as e:
                    logger.debug(f"Failed to parse date with format {date_format}: {e}")
                    continue
        
        logger.debug("Could not parse date")
        return None

    def calculate_years_experience(self, dates):
        """Calculate total years of experience from date ranges."""
        if not dates:
            logger.debug("No dates provided for experience calculation")
            return 0
            
        logger.debug(f"Calculating experience from dates: {dates}")
        total_years = 0
        current_date = datetime.now()
        
        for start_date, end_date in dates:
            if not start_date:
                logger.debug(f"Skipping date range due to missing start date: {start_date} - {end_date}")
                continue
                
            # Use current date if end_date is None (current position)
            end = end_date if end_date else current_date
            
            # Calculate years
            years = (end.year - start_date.year) + (end.month - start_date.month) / 12
            total_years += max(0, years)  # Ensure non-negative
            
            logger.debug(f"Date range {start_date} - {end}: {years:.1f} years")
            
        logger.debug(f"Total years experience: {total_years:.1f}")
        return round(total_years, 1)

    def _extract_name(self, text: str) -> str:
        """Extract name from text."""
        # Look for name at the start of the document
        lines = text.split('\n')
        for line in lines[:3]:  # Check first 3 lines
            # Look for capitalized words that could be a name
            words = line.strip().split()
            if len(words) >= 2 and all(w[0].isupper() for w in words if w):
                # Validate it's not a header or title
                if not any(x in line.lower() for x in ['resume', 'cv', 'curriculum vitae']):
                    return line.strip()
        return ""

    def _extract_role(self, text: str) -> str:
        """Extract current role from text."""
        # Common job titles and levels
        job_titles = [
            r'(?:senior|lead|principal|staff|chief|head|director|vp|manager|specialist|analyst|coordinator|consultant|generalist)',
            r'(?:software|systems|data|product|project|program|business|marketing|sales|hr|human\s+resources|operations|finance|recruitment)'
        ]
        
        # Look for role in professional summary or at the start
        summary_patterns = [
            # Standard role declarations
            r'(?i)(?:^|\n)(?:I am|Currently|Now|Presently)?\s*(?:a|an)?\s*([A-Z][a-zA-Z\s]+(?:' + '|'.join(job_titles) + r')(?:\s+at|\s+with|\s+for|\s+in|\s*\n|\s*$|\.))',
            r'(?i)current\s+(?:role|position|title):\s*([^\n]+)',
            r'(?i)(?:^|\n)experienced\s+([A-Z][a-zA-Z\s]+(?:' + '|'.join(job_titles) + r')(?:\s+with|\s+having|\.))',
            r'(?i)(?:^|\n)([A-Z][a-zA-Z\s]+(?:' + '|'.join(job_titles) + r'))\s+with\s+\d+\+?\s*years?',
            
            # Look for role at document start or after name
            r'(?i)(?:^|\n)([A-Z][a-zA-Z\s]+(?:' + '|'.join(job_titles) + r'))\s*(?:\n|$)',
            
            # Look for role in contact section
            r'(?i)(?:^|\n)(?:title|position|role):\s*([^\n]+)',
            
            # Look for role with company
            r'(?i)(?:^|\n)([A-Z][a-zA-Z\s]+(?:' + '|'.join(job_titles) + r'))\s*\|\s*[A-Z]',
            
            # Look for role in achievements
            r'(?i)(?:as\s+(?:a|an)\s+|(?:^|\n))([A-Z][a-zA-Z\s]+(?:' + '|'.join(job_titles) + r'))'
        ]
        
        # First try exact matches for HR roles
        hr_patterns = [
            r'(?i)(?:^|\n|\s)(?:Senior\s+)?(?:Human\s+Resources?|HR)\s+(?:Generalist|Manager|Specialist|Coordinator)(?:\s|$|\n)',
            r'(?i)(?:^|\n|\s)(?:Senior\s+)?(?:Human\s+Resources?|HR)\s+(?:Director|Consultant|Analyst)(?:\s|$|\n)'
        ]
        
        # Try to find the most recent role first
        experience_section = re.search(r'(?i)(?:experience|employment|work\s+history)[:\n]+(.*?)(?:\n\n|\Z)', text, re.DOTALL)
        if experience_section:
            experience_text = experience_section.group(1)
            # Look for role with date range
            date_role_pattern = r'(?i)(?:^|\n)(?:20\d{2}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*)\s*[-–]\s*(?:PRESENT|20\d{2})\s*\n([^\n|]+)'
            date_role_match = re.search(date_role_pattern, experience_text)
            if date_role_match:
                role = date_role_match.group(1).strip()
                if 2 <= len(role.split()) <= 6:
                    return role
        
        # Then try HR-specific patterns
        for pattern in hr_patterns:
            match = re.search(pattern, text)
            if match:
                role = match.group(0).strip()
                return role
        
        # Then try the general patterns
        for pattern in summary_patterns:
            match = re.search(pattern, text)
            if match:
                role = match.group(1).strip()
                # Remove trailing punctuation and company names
                role = re.sub(r'\s*(?:at|with|for|in)\s+.*$', '', role)
                role = role.strip(' .,')
                if 2 <= len(role.split()) <= 6:  # Reasonable length for a title
                    return role
        
        # Look for role in experience section
        if experience_section:
            lines = experience_section.group(1).split('\n')
            for line in lines[:3]:  # Check first few lines
                # Look for job title patterns
                for title in job_titles:
                    match = re.search(rf'(?i)([A-Z][a-zA-Z\s]*{title}[a-zA-Z\s]*)', line)
                    if match:
                        role = match.group(1).strip()
                        if 2 <= len(role.split()) <= 6:
                            return role
        
        # Try to extract from achievements or responsibilities
        achievements_section = re.search(r'(?i)(?:achievements|responsibilities|key\s+accomplishments)[:\n]+(.*?)(?:\n\n|\Z)', text, re.DOTALL)
        if achievements_section:
            lines = achievements_section.group(1).split('\n')
            for line in lines[:3]:
                for title in job_titles:
                    match = re.search(rf'(?i)(?:as\s+(?:a|an)\s+|(?:^|\n))([A-Z][a-zA-Z\s]*{title}[a-zA-Z\s]*)', line)
                    if match:
                        role = match.group(1).strip()
                        if 2 <= len(role.split()) <= 6:
                            return role
        
        return ""
    
    def _extract_text(self, doc: Document) -> str:
        """Extract all text from a docx document.
        
        Args:
            doc: Document to extract text from
            
        Returns:
            Extracted text
        """
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
        
        return '\n'.join(text)
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills from text."""
        skills = set()
        
        # Look for skills section
        skills_section = re.search(r'(?i)(?:skills|expertise|proficiencies)[:\n]+(.*?)(?:\n\n|\Z)', text, re.DOTALL)
        if skills_section:
            skill_text = skills_section.group(1)
            # Split by common delimiters
            skill_list = re.split(r'[•\n]', skill_text)
            
            for skill in skill_list:
                skill = skill.strip().lower()
                if skill:
                    # Check against our predefined skill sets
                    if (skill in self.TECHNICAL_SKILLS or
                        skill in self.BUSINESS_SKILLS or
                        skill in self.CREATIVE_SKILLS or
                        skill in self.HEALTHCARE_SKILLS or
                        skill in self.HR_SKILLS or
                        skill in self.SOFT_SKILLS or
                        skill in self.EDUCATION_SKILLS):
                        skills.add(skill)
                    # Check for compound skills
                    elif any(term in skill for term in ['management', 'development', 'analysis', 'planning']):
                        if not any(irr in skill for irr in self.IRRELEVANT_TERMS):
                            skills.add(skill)
        
        # Look for skills in text
        text_lower = text.lower()
        for skill_set in [self.TECHNICAL_SKILLS, self.BUSINESS_SKILLS, self.CREATIVE_SKILLS, self.HEALTHCARE_SKILLS, self.HR_SKILLS, self.SOFT_SKILLS, self.EDUCATION_SKILLS]:
            for skill in skill_set:
                if skill in text_lower:
                    skills.add(skill)
        
        # Look for additional technical terms
        skill_patterns = [
            r'(?i)(?:proficient|skilled|expertise|experienced)\s+(?:in|with)?\s+([^.,;]+)',
            r'(?i)knowledge\s+of\s+([^.,;]+)',
            r'(?i)experience\s+(?:in|with)\s+([^.,;]+)'
        ]
        
        for pattern in skill_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                skill = match.group(1).strip().lower()
                if (not any(irr in skill for irr in self.IRRELEVANT_TERMS) and
                    len(skill.split()) <= 3):
                    skills.add(skill)
        
        # Convert to list and sort by relevance
        skills_list = list(skills)
        skills_list.sort(key=lambda x: (
            x in self.TECHNICAL_SKILLS,  # Technical skills first
            x in self.HEALTHCARE_SKILLS,  # Healthcare skills second
            x in self.HR_SKILLS,  # HR skills third
            x in self.SOFT_SKILLS,  # Soft skills fourth
            x in self.EDUCATION_SKILLS,  # Education skills fifth
            x in self.CREATIVE_SKILLS,  # Creative skills sixth
            x in self.BUSINESS_SKILLS,  # Business skills seventh
            len(x)  # Shorter skills before longer ones
        ), reverse=True)
        
        return skills_list[:10]  # Return top 10 most relevant skills

    def _extract_companies(self, text: str) -> List[str]:
        """Extract companies from text."""
        companies = []
        
        # Look for company sections
        company_patterns = [
            r'(?i)company:\s*([^\n]+)',
            r'(?i)employer:\s*([^\n]+)',
            r'(?i)organization:\s*([^\n]+)',
            r'\b[A-Z][a-zA-Z\s&]+(?:Inc\.|LLC|Ltd\.|Corp\.|Corporation|Company)\b',
            r'(?i)(?:^|\n)(?:at|with|for)\s+([A-Z][a-zA-Z\s&]+)(?:\s+as|\s+in|\s+from|\n)',
        ]
        
        for pattern in company_patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                company = match.group(1).strip() if len(match.groups()) > 0 else match.group(0)
                if company and company not in companies:
                    # Validate it's a company name
                    if not any(x in company.lower() for x in ['resume', 'cv', 'summary', 'profile', 'experience']):
                        companies.append(company)
        
        return companies[:3]  # Return top 3 companies
    
    def _extract_years_experience(self, text: str) -> float:
        """Extract years of experience from text."""
        # Look for explicit mentions of years
        year_patterns = [
            r'(?i)(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?experience',
            r'(?i)experience:\s*(\d+)\+?\s*(?:years?|yrs?)',
            r'(?i)(?:^|\n)(\d+)\+?\s*(?:years?|yrs?)\s+(?:in|of|as)',
        ]
        
        for pattern in year_patterns:
            match = re.search(pattern, text)
            if match:
                try:
                    return float(match.group(1))
                except ValueError:
                    continue
        
        # Calculate from employment dates
        date_pattern = r'(?i)(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}'
        dates = re.findall(date_pattern, text)
        if len(dates) >= 2:
            try:
                dates = [datetime.strptime(d, '%B %Y') for d in dates]
                years = (max(dates) - min(dates)).days / 365.25
                return round(years, 1)
            except ValueError:
                pass
        
        return 0.0
    
    def _extract_achievements(self, text: str) -> List[Dict[str, Any]]:
        """Extract achievements with sophisticated metrics and impact analysis."""
        achievements = []
        
        # Achievement indicators with impact levels
        indicators = {
            'leadership': {
                'high': [
                    r'led\s+(?:global|enterprise|company-wide)',
                    r'directed\s+(?:strategic|major|key)',
                    r'spearheaded\s+(?:transformation|initiative)',
                    r'transformed\s+(?:organization|department)',
                    r'established\s+(?:new|innovative)\s+(?:division|department)'
                ],
                'medium': [
                    r'managed\s+(?:team|project|program)',
                    r'supervised\s+(?:staff|employees)',
                    r'coordinated\s+(?:efforts|activities)',
                    r'guided\s+(?:implementation|development)',
                    r'mentored\s+(?:team|staff|employees)'
                ],
                'low': [
                    r'assisted\s+(?:in|with)',
                    r'supported\s+(?:team|project)',
                    r'participated\s+(?:in|as)',
                    r'contributed\s+to',
                    r'helped\s+(?:with|in)'
                ]
            },
            'growth': {
                'high': [
                    r'doubled|tripled|quadrupled',
                    r'increased\s+(?:\d+(?:\.\d+)?x|\d{3,}%)',
                    r'grew\s+(?:revenue|sales|profit)\s+by\s+(?:\d{3,}%|\$\d+M)',
                    r'expanded\s+(?:globally|internationally)',
                    r'scaled\s+(?:operations|business)\s+(?:\d+x|\d{3,}%)'
                ],
                'medium': [
                    r'increased\s+by\s+(?:\d{2,3}%|\$\d+K)',
                    r'improved\s+(?:performance|efficiency)\s+by\s+\d{2,3}%',
                    r'enhanced\s+(?:productivity|output)\s+by\s+\d{2,3}%',
                    r'boosted\s+(?:sales|revenue)\s+by\s+\d{2,3}%',
                    r'grew\s+(?:team|department)\s+by\s+\d{2,3}%'
                ],
                'low': [
                    r'increased\s+by\s+(?:\d{1,2}%|\$\d+)',
                    r'improved\s+(?:slightly|marginally)',
                    r'enhanced\s+(?:somewhat|partially)',
                    r'contributed\s+to\s+growth',
                    r'supported\s+growth\s+initiatives'
                ]
            },
            'innovation': {
                'high': [
                    r'pioneered\s+(?:revolutionary|groundbreaking|first-ever)',
                    r'invented\s+(?:new|novel|innovative)',
                    r'developed\s+(?:patent|proprietary)',
                    r'created\s+(?:revolutionary|breakthrough)',
                    r'designed\s+(?:award-winning|innovative)'
                ],
                'medium': [
                    r'implemented\s+(?:new|improved)',
                    r'redesigned\s+(?:process|system)',
                    r'modernized\s+(?:approach|method)',
                    r'enhanced\s+(?:technology|system)',
                    r'upgraded\s+(?:platform|infrastructure)'
                ],
                'low': [
                    r'assisted\s+(?:in|with)\s+development',
                    r'supported\s+(?:implementation|rollout)',
                    r'helped\s+(?:develop|create)',
                    r'participated\s+in\s+(?:development|design)',
                    r'contributed\s+to\s+(?:project|initiative)'
                ]
            },
            'efficiency': {
                'high': [
                    r'reduced\s+(?:costs|expenses)\s+by\s+(?:\d{3,}%|\$\d+M)',
                    r'automated\s+(?:\d{2,}|multiple)\s+processes',
                    r'eliminated\s+(?:\d{2,}%|major)\s+(?:waste|redundancy)',
                    r'optimized\s+(?:enterprise|company-wide)\s+operations',
                    r'streamlined\s+(?:critical|key)\s+(?:processes|operations)'
                ],
                'medium': [
                    r'reduced\s+(?:time|costs)\s+by\s+\d{2,3}%',
                    r'improved\s+efficiency\s+by\s+\d{2,3}%',
                    r'automated\s+(?:process|workflow)',
                    r'streamlined\s+(?:operations|procedures)',
                    r'optimized\s+(?:workflow|process)'
                ],
                'low': [
                    r'reduced\s+(?:time|costs)\s+by\s+\d{1,2}%',
                    r'improved\s+(?:slightly|somewhat)',
                    r'helped\s+streamline',
                    r'assisted\s+with\s+optimization',
                    r'supported\s+efficiency\s+initiatives'
                ]
            },
            'impact': {
                'high': [
                    r'generated\s+(?:\$\d+M|\d{3,}%)',
                    r'saved\s+(?:\$\d+M|\d{3,}%)',
                    r'impacted\s+(?:company-wide|enterprise)',
                    r'transformed\s+(?:industry|market)',
                    r'revolutionized\s+(?:approach|method)'
                ],
                'medium': [
                    r'generated\s+(?:\$\d+K|\d{2,3}%)',
                    r'saved\s+(?:\$\d+K|\d{2,3}%)',
                    r'improved\s+(?:key|important)',
                    r'enhanced\s+(?:significant|substantial)',
                    r'delivered\s+(?:significant|substantial)'
                ],
                'low': [
                    r'generated\s+(?:\$\d+|\d{1,2}%)',
                    r'saved\s+(?:\$\d+|\d{1,2}%)',
                    r'improved\s+(?:minor|small)',
                    r'contributed\s+to',
                    r'supported\s+(?:efforts|initiatives)'
                ]
            }
        }
        
        # Metric patterns with sophistication levels
        metric_patterns = {
            'percentage': {
                'high': r'(\d{3,}(?:\.\d+)?%|\d{3,}(?:\.\d+)?\spercent)',
                'medium': r'(\d{2}(?:\.\d+)?%|\d{2}(?:\.\d+)?\spercent)',
                'low': r'(\d{1}(?:\.\d+)?%|\d{1}(?:\.\d+)?\spercent)'
            },
            'money': {
                'high': r'[\$£€](\d+(?:\.\d+)?[mM]|\d{7,})',
                'medium': r'[\$£€](\d+(?:\.\d+)?[kK]|\d{4,6})',
                'low': r'[\$£€](\d{1,3}(?:,\d{3})*(?:\.\d{2})?)'
            },
            'scale': {
                'high': r'(\d{2,}x|\d{2,}\stimes)',
                'medium': r'(\d+\.\d+x|\d+\.\d+\stimes)',
                'low': r'(\d+(?:\.\d+)?x|\d+(?:\.\d+)?\stimes)'
            },
            'time': {
                'high': r'(\d+\+?\s+years?)',
                'medium': r'(\d+\s+(?:month|quarter)s?)',
                'low': r'(\d+\s+(?:day|week)s?)'
            },
            'quantity': {
                'high': r'(\d{5,}(?:\+|\s*\+)?)',
                'medium': r'(\d{3,4}(?:\+|\s*\+)?)',
                'low': r'(\d{1,2}(?:\+|\s*\+)?)'
            }
        }
        
        # Impact keywords with weights
        impact_keywords = {
            'high': {
                'global': 5, 'enterprise': 5, 'company-wide': 5,
                'revolutionary': 5, 'breakthrough': 5, 'transformative': 5,
                'first-ever': 5, 'groundbreaking': 5, 'pioneering': 5
            },
            'medium': {
                'significant': 3, 'substantial': 3, 'important': 3,
                'valuable': 3, 'key': 3, 'major': 3, 'critical': 3,
                'essential': 3, 'strategic': 3
            },
            'low': {
                'improved': 1, 'enhanced': 1, 'supported': 1,
                'assisted': 1, 'helped': 1, 'contributed': 1
            }
        }
        
        # Extract sentences that might contain achievements
        sentences = re.split(r'[.!?]+', text)
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence or len(sentence.split()) < 4 or sentence.isupper():
                continue
                
            # Skip sentences with contact information
            if self._is_contact_info(sentence):
                continue
                
            # Skip sentences with dates or role titles
            if self._is_date_or_role(sentence):
                continue
            
            # Skip sentences that look like headers or section titles
            if all(word[0].isupper() for word in sentence.split() if word):
                continue
            
            # Skip education-related sentences
            if any(edu_term in sentence.lower() for edu_term in ['bachelor', 'master', 'phd', 'degree', 'university', 'college']):
                continue
            
            # Skip non-achievement sentences
            if not any(word in sentence.lower() for word in [
                'led', 'managed', 'created', 'developed', 'implemented',
                'improved', 'increased', 'reduced', 'achieved', 'delivered',
                'launched', 'designed', 'established', 'streamlined', 'optimized'
            ]):
                continue
            
            achievement_data = {
                'text': sentence,
                'categories': {},
                'metrics': {},
                'impact_level': 'low',
                'score': 0
            }
            
            # Score achievement indicators
            for category, levels in indicators.items():
                category_score = 0
                for level, patterns in levels.items():
                    for pattern in patterns:
                        if re.search(rf'\b{pattern}\b', sentence.lower()):
                            level_scores = {'high': 5, 'medium': 3, 'low': 1}
                            category_score = max(category_score, level_scores[level])
                
                if category_score > 0:
                    achievement_data['categories'][category] = category_score
                    achievement_data['score'] += category_score
            
            # Score metrics
            for metric_type, levels in metric_patterns.items():
                metric_matches = {}
                for level, pattern in levels.items():
                    matches = re.findall(pattern, sentence, re.IGNORECASE)
                    if matches:
                        level_scores = {'high': 5, 'medium': 3, 'low': 1}
                        metric_matches[level] = matches
                        achievement_data['score'] += level_scores[level] * len(matches)
                
                if metric_matches:
                    achievement_data['metrics'][metric_type] = metric_matches
            
            # Score impact keywords
            impact_score = 0
            for level, keywords in impact_keywords.items():
                for keyword, weight in keywords.items():
                    if keyword in sentence.lower():
                        impact_score = max(impact_score, weight)
            
            if impact_score > 0:
                achievement_data['score'] += impact_score
            
            # Determine overall impact level
            if achievement_data['score'] >= 15:
                achievement_data['impact_level'] = 'high'
            elif achievement_data['score'] >= 8:
                achievement_data['impact_level'] = 'medium'
            
            # Time-based scoring (recent achievements score higher)
            time_indicators = {
                'current': 5,
                'recently': 4,
                'this year': 3,
                'last year': 2,
                'previously': 1
            }
            for indicator, score in time_indicators.items():
                if indicator in sentence.lower():
                    achievement_data['score'] += score
                    break
            
            # Only include if it has some indicators of being an achievement
            if achievement_data['score'] > 0:
                # Clean up the achievement text
                cleaned_text = achievement_data['text']
                cleaned_text = re.sub(r'\s*\|\s*', ' at ', cleaned_text)  # Replace | with 'at'
                cleaned_text = re.sub(r'\s*[-–]\s*', ' to ', cleaned_text)  # Replace - with 'to'
                cleaned_text = re.sub(r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\b', '', cleaned_text)  # Remove dates
                cleaned_text = re.sub(r'\b\d{4}\s*[-–]\s*(?:Present|Current|Now|\d{4})\b', '', cleaned_text)  # Remove date ranges
                cleaned_text = re.sub(r'\s{2,}', ' ', cleaned_text)  # Remove extra spaces
                achievement_data['text'] = cleaned_text.strip()
                
                achievements.append(achievement_data)
        
        # Sort by score and return top achievements
        achievements.sort(key=lambda x: x['score'], reverse=True)
        return achievements[:5]
    
    def _format_achievement(self, achievement: Dict[str, Any]) -> str:
        """Format achievement data into a readable string with impact context."""
        text = achievement['text']
        
        # Add impact-based context
        impact_context = {
            'high': {
                'leadership': 'Through transformative leadership',
                'growth': 'Delivering exceptional growth',
                'efficiency': 'Dramatically improving efficiency',
                'innovation': 'Pioneering innovative solutions',
                'impact': 'Creating transformative impact'
            },
            'medium': {
                'leadership': 'Through effective leadership',
                'growth': 'Driving significant growth',
                'efficiency': 'Substantially improving efficiency',
                'innovation': 'Implementing innovative solutions',
                'impact': 'Delivering meaningful impact'
            },
            'low': {
                'leadership': 'Supporting leadership initiatives',
                'growth': 'Contributing to growth',
                'efficiency': 'Improving efficiency',
                'innovation': 'Supporting innovation',
                'impact': 'Making positive impact'
            }
        }
        
        # Get primary category (highest scoring)
        if achievement['categories']:
            primary_category = max(achievement['categories'].items(), key=lambda x: x[1])[0]
            impact_level = achievement['impact_level']
            context = impact_context[impact_level][primary_category]
            text = f"{context} - {text}"
        
        # Highlight metrics based on their level
        if achievement['metrics']:
            for metric_type, levels in achievement['metrics'].items():
                for level, values in levels.items():
                    for value in values:
                        highlight = {
                            'high': f"**{value}**",
                            'medium': f"*{value}*",
                            'low': value
                        }
                        text = text.replace(str(value), highlight[level])
        
        return text
    
    def _parse_achievements(self, doc: Document) -> List[str]:
        """Parse achievements from the document.
        
        Args:
            doc: Document to parse
            
        Returns:
            List of parsed achievements
        """
        text = self._extract_text(doc)
        achievements = self._extract_achievements(text)
        
        if not achievements:
            logger.warning("No achievements found in document")
            return []
        
        return [self._format_achievement(achievement) for achievement in achievements]

    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from text."""
        contact_info = {
            'email': '',
            'phone': ''
        }
        
        # Extract email
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            contact_info['email'] = email_match.group(0)
        
        # Extract phone with various formats
        phone_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # Standard US format
            r'\(\d{3}\)\s*\d{3}[-.]?\d{4}',     # (123) 456-7890
            r'\+\d{1,2}\s*\d{3}[-.]?\d{3}[-.]?\d{4}'  # International
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                # Clean up phone format
                phone = re.sub(r'[^\d]', '', phone_match.group(0))
                if len(phone) >= 10:  # Ensure it's a valid length
                    contact_info['phone'] = f"({phone[-10:-7]}) {phone[-7:-4]}-{phone[-4:]}"
                break
        
        return contact_info

    def _is_contact_info(self, text: str) -> bool:
        """Check if text contains contact information."""
        contact_patterns = [
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',  # Phone numbers
            r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Email
            r'\b\d{5}(?:[-\s]\d{4})?\b',  # ZIP codes
            r'\b\d+\s+[A-Za-z\s]+(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Circle|Cir|Way|Place|Pl)\b',  # Addresses
            r'\b(?:http[s]?://)?(?:www\.)?[a-zA-Z0-9-]+(?:\.[a-zA-Z]{2,})+\b',  # URLs
            r'\b(?:linkedin\.com|github\.com|twitter\.com)/[\w-]+\b'  # Social media
        ]
        
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in contact_patterns)

    def _is_date_or_role(self, text: str) -> bool:
        """Check if text contains dates or role titles."""
        date_role_patterns = [
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December|Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\b',
            r'\b\d{4}\s*[-–]\s*(?:Present|Current|Now|\d{4})\b',
            r'\b(?:Senior|Junior|Lead|Principal|Associate|Assistant)\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b',
            r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Manager|Director|Specialist|Analyst|Engineer|Developer|Consultant)\b',
            r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+\|\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b'  # Role | Company format
        ]
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in date_role_patterns)

    def _process_content(self, text_content: List[str]) -> None:
        """Process the extracted text content.
        
        Args:
            text_content: List of text lines from the document
        """
        try:
            # Join all text for processing
            full_text = '\n'.join(text_content)
            logger.debug("Processing text content:")
            logger.debug("-" * 40)
            logger.debug(full_text)
            logger.debug("-" * 40)
            
            # Extract basic information
            self.resume_data['name'] = self._extract_name(full_text)
            logger.debug(f"Extracted name: {self.resume_data['name']}")
            
            self.resume_data['current_role'] = self._extract_role(full_text)
            logger.debug(f"Extracted role: {self.resume_data['current_role']}")
            
            self.resume_data['skills'] = self._extract_skills(full_text)
            logger.debug(f"Extracted skills: {self.resume_data['skills']}")
            
            self.resume_data['companies'] = self._extract_companies(full_text)
            logger.debug(f"Extracted companies: {self.resume_data['companies']}")
            
            self.resume_data['years_experience'] = self._extract_years_experience(full_text)
            logger.debug(f"Extracted years: {self.resume_data['years_experience']}")
            
            # Extract contact information
            contact_info = self._extract_contact_info(full_text)
            if contact_info:
                self.resume_data['contact_info'] = contact_info
                self.resume_data['email'] = contact_info.get('email', '')
            logger.debug(f"Extracted contact: {contact_info}")
            
            # Process achievements
            doc = Document(self.file_path)
            achievements = self._parse_achievements(doc)
            if achievements:
                self.resume_data['achievements'] = achievements
            logger.debug(f"Extracted achievements: {achievements}")
            
            # Extract education information
            education_info = []
            education_section = False
            education_entry = {}
            
            for line in text_content:
                line = line.strip()
                if not line:
                    continue
                    
                # Look for education section
                if re.search(r'education|academic|qualification', line.lower()):
                    education_section = True
                    continue
                
                if education_section:
                    # Look for degree information
                    degree_match = re.search(r"(?:Bachelor's|Master's|PhD|B\.[A-Z]|M\.[A-Z]|Ph\.D)\s+(?:of|in|degree in)?\s+([^\n]+)", line)
                    if degree_match:
                        education_entry = {'degree': degree_match.group(0)}
                        education_info.append(education_entry)
                        education_entry = {}
            
            if education_info:
                self.resume_data['education'] = education_info
            logger.debug(f"Extracted education: {education_info}")
            
        except Exception as e:
            logger.error(f"Error processing content: {str(e)}")
            raise

# Example usage
if __name__ == "__main__":
    file_path = 'src/templates/ATS classic HR resume.docx'
    parser = ATSParser(file_path)
    parsed_data = parser.parse_docx_to_json()
    print(json.dumps(parsed_data, indent=4))
